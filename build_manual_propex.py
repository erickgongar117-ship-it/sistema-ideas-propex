from pathlib import Path
from math import atan2, cos, sin, pi

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "manual-propex-assets"
OUTPUT = ROOT / "Manual_de_Usuario_e_Instructivo_PROpEx.docx"

RED = "D61F3C"
RED_DARK = "9F1239"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F4F6F8"
PALE_RED = "FFF0F3"
GREEN = "14835F"
BLUE = "176FC1"
AMBER = "B76800"
PURPLE = "7C3AED"
WHITE = "FFFFFF"
BORDER = "D9DEE7"


def rgb(value):
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 10),
        ("Heading 1", 16, RED, 18, 10),
        ("Heading 2", 13, RED_DARK, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("PROpEx  |  Manual de usuario e instructivo de trabajo")
    set_run(r, 8.5, bold=True, color=MUTED)
    p_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = fp.add_run("Proboca · Uso interno · Versión 1.0 · julio 2026")
    set_run(left, 8.5, color=MUTED)


def paragraph(doc, text="", size=11, bold=False, color=INK, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    set_run(p.add_run(text), size, bold, color, italic)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text.upper()), 9, True, RED)
    return p


def add_callout(doc, label, text, fill=PALE_RED, accent=RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + ": "), 10.5, True, accent)
    set_run(p.add_run(text), 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, steps):
    for title, detail in steps:
        p = doc.add_paragraph(style="List Number")
        set_run(p.add_run(title + ". "), 11, True, INK)
        set_run(p.add_run(detail), 11, False, INK)


def add_role_table(doc):
    rows = [
        ("Colaborador / operador", "Captura por QR", "Problema, propuesta, beneficio, impacto y evidencia"),
        ("Supervisor", "Revisión inicial", "Aprobar, rechazar o solicitar información; confirmar impactos"),
        ("Calidad / Inocuidad", "Validación especializada", "Riesgo a producto, proceso, higiene y controles"),
        ("Seguridad Industrial", "Validación especializada", "Riesgos de seguridad, ergonomía y cumplimiento"),
        ("Mantenimiento", "Validación especializada", "Factibilidad técnica, recursos y restricciones"),
        ("Mejora Continua", "Gobierno del flujo", "Clasificar, priorizar, asignar, cerrar y otorgar puntos"),
        ("Responsable de implementación", "Ejecución", "Avances, fechas, evidencia antes/después y resultado"),
        ("Dirección", "Decisiones escaladas", "Resolver desbordes de costo, autoridad o capacidad"),
        ("Administrador", "Configuración", "Usuarios, accesos, áreas, supervisores, reglas y auditoría"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2300, 2250, 4810])
    headers = ["Perfil", "Momento", "Responsabilidad principal"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], RED)
        p = table.rows[0].cells[i].paragraphs[0]
        set_run(p.add_run(h), 9.5, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        if idx % 2:
            for c in cells:
                set_cell_shading(c, LIGHT)
        for col, text in enumerate(row):
            p = cells[col].paragraphs[0]
            set_run(p.add_run(text), 9.5, col == 0, INK)
    set_table_geometry(table, [2300, 2250, 4810])
    return table


def set_picture_alt(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_figure(doc, filename, caption, width=6.72):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    set_picture_alt(shape, caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    set_run(cp.add_run(caption), 9, False, MUTED, True)


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size=size)
    return ImageFont.load_default()


def wrap(draw, text, fnt, max_width):
    words = text.split()
    lines, current = [], ""
    for word in words:
        trial = (current + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(draw, xy, fill, title, detail="", title_color="#FFFFFF", detail_color="#FFFFFF", radius=24):
    draw.rounded_rectangle(xy, radius=radius, fill=fill)
    x1, y1, x2, y2 = xy
    tf = font(30, True)
    df = font(23, False)
    y = y1 + 20
    for line in wrap(draw, title, tf, x2 - x1 - 40):
        bbox = draw.textbbox((0, 0), line, font=tf)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=tf, fill=title_color)
        y += 36
    if detail:
        y += 4
        for line in wrap(draw, detail, df, x2 - x1 - 48):
            bbox = draw.textbbox((0, 0), line, font=df)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=df, fill=detail_color)
            y += 29


def arrow(draw, start, end, color="#667085", width=7):
    draw.line([start, end], fill=color, width=width)
    angle = atan2(end[1] - start[1], end[0] - start[0])
    length = 24
    points = [
        end,
        (end[0] - length * cos(angle - pi / 6), end[1] - length * sin(angle - pi / 6)),
        (end[0] - length * cos(angle + pi / 6), end[1] - length * sin(angle + pi / 6)),
    ]
    draw.polygon(points, fill=color)


def make_process_flow():
    img = Image.new("RGB", (1800, 2700), "white")
    d = ImageDraw.Draw(img)
    d.text((90, 60), "Flujo operativo PROpEx", font=font(54, True), fill="#1F2937")
    d.text((90, 130), "De la idea en piso al cierre y aprendizaje", font=font(30), fill="#667085")

    lane_x = 90
    lane_w = 330
    process_x1, process_x2 = 500, 1710
    lanes = [
        (220, 560, "OPERADOR /\nCOLABORADOR", "#D61F3C"),
        (590, 860, "SUPERVISOR", "#14835F"),
        (890, 1300, "ÁREAS DE\nVALIDACIÓN", "#176FC1"),
        (1330, 1760, "MEJORA CONTINUA", "#7C3AED"),
        (1790, 2170, "RESPONSABLE", "#B76800"),
        (2200, 2590, "SISTEMA /\nDIRECCIÓN", "#344054"),
    ]
    for y1, y2, label, color in lanes:
        d.rounded_rectangle((lane_x, y1, lane_x + lane_w, y2), radius=25, fill=color)
        yy = (y1 + y2) // 2 - 34
        for line in label.split("\n"):
            bbox = d.textbbox((0, 0), line, font=font(27, True))
            d.text((lane_x + (lane_w - (bbox[2] - bbox[0])) / 2, yy), line, font=font(27, True), fill="white")
            yy += 36

    rounded_box(d, (500, 245, 1010, 395), "#D61F3C", "1. Escanear QR", "Seleccionar el área correcta")
    rounded_box(d, (1120, 245, 1710, 395), "#D61F3C", "2. Registrar idea", "Problema + propuesta + beneficio + evidencia")
    arrow(d, (1010, 320), (1120, 320))
    arrow(d, (1415, 395), (1415, 610))

    rounded_box(d, (880, 620, 1450, 800), "#14835F", "3. Revisar", "Aprobar · pedir información · rechazar")
    rounded_box(d, (1490, 620, 1710, 800), "#FDECEC", "¿Aprobar?", "", title_color="#9F1239")
    arrow(d, (1450, 710), (1490, 710))
    d.text((1535, 820), "No: regresa o termina", font=font(21, True), fill="#9F1239")
    arrow(d, (1600, 800), (1600, 965), color="#9F1239")

    d.text((500, 900), "VALIDACIONES EN PARALELO SEGÚN IMPACTO", font=font(25, True), fill="#176FC1")
    rounded_box(d, (500, 965, 850, 1170), "#176FC1", "Calidad / Inocuidad", "Riesgo a producto y proceso")
    rounded_box(d, (915, 965, 1270, 1170), "#50606D", "Seguridad", "Riesgo, ergonomía y controles")
    rounded_box(d, (1335, 965, 1710, 1170), "#176FC1", "Mantenimiento", "Factibilidad y recursos")
    arrow(d, (1100, 800), (675, 965))
    arrow(d, (1165, 800), (1090, 965))
    arrow(d, (1230, 800), (1520, 965))
    d.text((500, 1200), "El sistema espera todas las validaciones obligatorias.", font=font(23), fill="#667085")
    arrow(d, (1090, 1170), (1090, 1370))

    rounded_box(d, (500, 1380, 1030, 1575), "#7C3AED", "4. Clasificar y priorizar", "Tipo · impacto · urgencia · valor")
    rounded_box(d, (1120, 1380, 1710, 1575), "#7C3AED", "5. Planificar", "Responsable + fecha compromiso + recursos")
    arrow(d, (1030, 1475), (1120, 1475))
    arrow(d, (1415, 1575), (1415, 1830))

    rounded_box(d, (500, 1840, 1050, 2045), "#B76800", "6. Implementar", "Registrar avances y bloqueos")
    rounded_box(d, (1140, 1840, 1710, 2045), "#B76800", "7. Evidenciar", "Antes / después + resultado")
    arrow(d, (1050, 1940), (1140, 1940))
    arrow(d, (1415, 2045), (1415, 2230))

    rounded_box(d, (500, 2240, 1030, 2455), "#344054", "8. Validar cierre", "Mejora Continua confirma el resultado")
    rounded_box(d, (1120, 2240, 1710, 2455), "#14835F", "9. Cerrar y aprender", "Puntos + notificación + indicadores + auditoría")
    arrow(d, (1030, 2345), (1120, 2345))

    d.rounded_rectangle((500, 2520, 1710, 2640), radius=22, fill="#F4F6F8", outline="#D9DEE7", width=3)
    d.text((540, 2545), "Ruta especial: si requiere autoridad, presupuesto o capacidad adicional,", font=font(23, True), fill="#1F2937")
    d.text((540, 2580), "se genera un desborde a Dirección y la decisión regresa al plan de acción.", font=font(23), fill="#667085")
    path = ASSETS / "flujo-operativo-propex.png"
    img.save(path, quality=95)
    return path


def make_information_flow():
    img = Image.new("RGB", (1800, 1450), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), "Flujo de la información", font=font(52, True), fill="#1F2937")
    d.text((80, 122), "Un folio único alimenta toda la trazabilidad", font=font(29), fill="#667085")

    headers = [(80, "ENTRADAS", "#D61F3C"), (640, "REGISTRO MAESTRO", "#344054"), (1230, "SALIDAS", "#14835F")]
    for x, label, color in headers:
        d.rounded_rectangle((x, 215, x + 480, 290), radius=18, fill=color)
        bbox = d.textbbox((0, 0), label, font=font(26, True))
        d.text((x + (480 - (bbox[2] - bbox[0])) / 2, 235), label, font=font(26, True), fill="white")

    inputs = [
        ("Captura en piso", "Problema, propuesta, beneficio"),
        ("Decisiones", "Supervisor y validadores"),
        ("Ejecución", "Avances, fechas y evidencia"),
        ("Cierre", "Resultado y puntos"),
    ]
    y = 345
    for title, detail in inputs:
        rounded_box(d, (80, y, 560, y + 185), "#FFF0F3", title, detail, title_color="#9F1239", detail_color="#667085")
        arrow(d, (560, y + 92), (640, 720), color="#D61F3C", width=5)
        y += 235

    d.rounded_rectangle((640, 355, 1120, 1090), radius=32, fill="#344054")
    d.text((740, 405), "FOLIO PROpEx", font=font(38, True), fill="white")
    d.text((765, 460), "IM-000000", font=font(31, True), fill="#FECACA")
    master_items = ["Datos de la idea", "Estatus y responsable", "Validaciones", "Comentarios", "Evidencias", "Puntos", "Bitácora de auditoría", "Notificaciones"]
    yy = 545
    for item in master_items:
        d.ellipse((695, yy + 8, 715, yy + 28), fill="#FECACA")
        d.text((735, yy), item, font=font(24), fill="white")
        yy += 63

    outputs = [
        ("Trabajo diario", "Bandejas y pendientes por rol"),
        ("Control", "Detalle del folio y Kanban"),
        ("Dirección", "Tablero, desbordes e indicadores"),
        ("Respaldo", "Excel, notificaciones y auditoría"),
    ]
    y = 345
    for title, detail in outputs:
        arrow(d, (1120, 720), (1230, y + 92), color="#14835F", width=5)
        rounded_box(d, (1230, y, 1710, y + 185), "#E9F6F0", title, detail, title_color="#0F684D", detail_color="#667085")
        y += 235

    d.rounded_rectangle((80, 1250, 1710, 1370), radius=22, fill="#F4F6F8", outline="#D9DEE7", width=3)
    d.text((120, 1275), "Regla de control: no crear seguimientos paralelos fuera del folio.", font=font(27, True), fill="#1F2937")
    d.text((120, 1320), "Toda decisión, avance y evidencia debe quedar registrada en PROpEx.", font=font(25), fill="#667085")
    path = ASSETS / "flujo-informacion-propex.png"
    img.save(path, quality=95)
    return path


def add_section_page(doc, kicker, title, intro):
    doc.add_page_break()
    add_kicker(doc, kicker)
    doc.add_heading(title, level=1)
    paragraph(doc, intro, 11.5, color=MUTED, after=10)


def build():
    ASSETS.mkdir(exist_ok=True)
    make_process_flow()
    make_information_flow()

    doc = Document()
    configure_document(doc)

    # Cover: editorial-cover pattern with restrained PROpEx branding.
    paragraph(doc, "PROBOCA", 11, True, RED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    if (ROOT / "public/brand/proboca-logo.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_shape = p.add_run().add_picture(str(ROOT / "public/brand/proboca-logo.png"), width=Inches(2.1))
        set_picture_alt(logo_shape, "Logotipo de Proboca")
        p.paragraph_format.space_after = Pt(28)
    paragraph(doc, "MANUAL DE USUARIO", 28, True, INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    paragraph(doc, "E INSTRUCTIVO DE TRABAJO", 22, True, RED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    paragraph(doc, "Sistema PROpEx · Ideas de mejora, proyectos Kaizen y recorridos Genba", 13, False, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_callout(doc, "Propósito", "Asegurar que cada usuario sepa qué información registrar, en qué momento hacerlo y cómo conservar la trazabilidad desde la idea en piso hasta el cierre.")
    paragraph(doc, "Dirigido a operadores, colaboradores, supervisores, áreas de soporte, responsables de implementación, Mejora Continua, Dirección y administradores.", 11, align=WD_ALIGN_PARAGRAPH.CENTER, after=38)
    paragraph(doc, "Versión 1.0  |  13 de julio de 2026  |  Uso interno", 10, True, MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    paragraph(doc, "Documento base para capacitación y validación operativa", 9.5, False, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_section_page(doc, "Antes de comenzar", "1. Cómo usar este manual", "Este documento combina una guía rápida por perfil con instrucciones paso a paso. Las capturas corresponden al entorno de demostración de PROpEx disponible al elaborar esta versión.")
    doc.add_heading("Objetivo", level=2)
    paragraph(doc, "Estandarizar el uso del sistema, evitar seguimientos paralelos y asegurar que las decisiones, responsables, fechas y evidencias queden disponibles para consulta y auditoría.")
    doc.add_heading("Ruta de lectura recomendada", level=2)
    add_bullets(doc, [
        "Operador o colaborador: secciones 2, 5 y 6.",
        "Supervisor y áreas de validación: secciones 2, 7 y 8.",
        "Mejora Continua y responsables: secciones 2, 9, 10 y 11.",
        "Dirección y administración: secciones 3, 4, 12 y 13.",
    ])
    add_callout(doc, "Regla de oro", "El folio es la fuente oficial. No sustituir el registro de PROpEx con mensajes, notas sueltas o archivos personales.", fill="FFF7E6", accent=AMBER)
    doc.add_heading("Alcance", level=2)
    paragraph(doc, "Incluye ideas de mejora, proyectos Kaizen, recorridos Genba, desbordes a Dirección, tableros, reportes, notificaciones y controles administrativos. Los nombres de campos pueden evolucionar; la responsabilidad sobre la calidad del dato se mantiene.")

    add_section_page(doc, "Gobierno del proceso", "2. Perfiles y responsabilidades", "Cada perfil ve información y acciones diferentes. La siguiente matriz indica quién debe registrar cada tipo de dato.")
    add_role_table(doc)
    doc.add_heading("Separación de responsabilidades", level=2)
    add_bullets(doc, [
        "Quien propone describe la situación; no necesita definir la solución técnica final.",
        "El supervisor valida pertinencia y contexto del área.",
        "Las áreas de soporte validan sólo el riesgo o factibilidad de su especialidad.",
        "Mejora Continua gobierna clasificación, prioridad, asignación y cierre.",
        "El responsable ejecuta y demuestra; no debe cerrar su propio resultado sin validación.",
    ])

    add_section_page(doc, "Visión integral", "3. Flujo operativo", "El proceso inicia en piso, abre validaciones según el impacto y termina cuando el resultado queda demostrado y comunicado.")
    add_figure(doc, "flujo-operativo-propex.png", "Figura 1. Flujo operativo PROpEx, incluyendo validaciones paralelas y ruta de desborde.", width=6.55)

    add_section_page(doc, "Trazabilidad", "4. Flujo de la información", "PROpEx conserva un registro maestro por folio. Las pantallas, indicadores, reportes y notificaciones deben derivarse de ese mismo registro.")
    add_figure(doc, "flujo-informacion-propex.png", "Figura 2. Entradas, registro maestro y salidas de información.", width=6.55)
    doc.add_heading("Reglas de calidad del dato", level=2)
    add_bullets(doc, [
        "Registrar hechos observables: qué sucede, dónde, cuándo y con qué frecuencia.",
        "Evitar datos personales innecesarios o información sensible en fotografías y comentarios.",
        "Usar comentarios para decisiones y aclaraciones; usar evidencia para demostrar condiciones o resultados.",
        "Actualizar responsable, fecha y estatus en cuanto cambien; no esperar al cierre.",
        "No eliminar la historia: corregir mediante una nueva actualización que explique el cambio.",
    ])

    add_section_page(doc, "Acceso", "5. Ingreso y navegación", "Los usuarios con rol ingresan con correo y contraseña. Los operadores pueden registrar una idea por QR sin iniciar sesión.")
    add_figure(doc, "01-inicio-sesion.png", "Figura 3. Pantalla de inicio de sesión para usuarios con rol.")
    doc.add_heading("Usuarios con cuenta", level=2)
    add_steps(doc, [
        ("Abrir PROpEx", "Ingresar a la dirección definida por la organización"),
        ("Capturar credenciales", "Escribir correo institucional y contraseña"),
        ("Entrar", "Seleccionar Entrar y confirmar que el nombre y rol sean correctos"),
        ("Ubicar el módulo", "Usar el menú lateral: Mi trabajo, Ideas, Kaizen, Genba, análisis o administración"),
        ("Cerrar sesión", "Seleccionar Salir cuando termine el trabajo o comparta el equipo"),
    ])
    add_callout(doc, "Seguridad", "No compartir contraseñas ni dejar abierta una sesión en equipos de uso común.", fill="FFF7E6", accent=AMBER)

    add_section_page(doc, "Operadores y colaboradores", "6. Registrar una idea por QR", "Este procedimiento está diseñado para ejecutarse desde un teléfono o dispositivo disponible en el área.")
    add_figure(doc, "02-captura-operador.png", "Figura 4. Formulario público de captura para el área P1.")
    doc.add_heading("Antes de capturar", level=2)
    add_bullets(doc, [
        "Confirmar que el QR corresponde al área donde se detectó la oportunidad.",
        "Tomar una fotografía sólo cuando aporte contexto y esté permitido por las reglas de planta.",
        "Separar el problema actual de la propuesta de mejora.",
    ])
    doc.add_heading("Procedimiento", level=2)
    add_steps(doc, [
        ("Escanear", "Abrir el QR del área y verificar el nombre mostrado en la cabecera"),
        ("Identificarse", "Registrar nombre; número de empleado y correo sólo cuando sean solicitados o útiles para seguimiento"),
        ("Seleccionar turno", "Elegir el turno en el que se detectó la situación"),
        ("Describir el problema", "Indicar la condición actual con hechos, ubicación y consecuencia"),
        ("Proponer", "Explicar qué podría cambiarse sin prometer una solución no validada"),
        ("Indicar beneficio", "Expresar el efecto esperado: seguridad, calidad, tiempo, costo, ergonomía u otro"),
        ("Adjuntar evidencia", "Agregar imagen cuando sea pertinente y revisar que no contenga información sensible"),
        ("Enviar", "Confirmar el registro y conservar el folio para futuras consultas"),
    ])
    add_callout(doc, "Ejemplo útil", "Problema: la manguera no alcanza la canaleta y quedan residuos. Propuesta: reubicar la toma o usar una manguera retráctil. Beneficio: limpieza completa en menos tiempo.")

    add_section_page(doc, "Trabajo diario", "7. Bandeja, búsqueda y detalle del folio", "La bandeja reúne las ideas y permite filtrar por folio, área o estatus. El detalle muestra la historia completa.")
    add_figure(doc, "04-bandeja-ideas.png", "Figura 5. Bandeja maestra de ideas con filtros y estados.")
    add_figure(doc, "11-detalle-idea.png", "Figura 6. Detalle del folio con avance por etapas y contenido de la oportunidad.")
    doc.add_heading("Cómo dar seguimiento", level=2)
    add_steps(doc, [
        ("Localizar", "Buscar por folio, texto del problema o colaborador; complementar con filtros"),
        ("Abrir", "Seleccionar Ver o Abrir para consultar el expediente"),
        ("Leer el avance", "Confirmar etapa actual, estatus, responsable y fechas"),
        ("Revisar historia", "Consultar validaciones, comentarios, evidencia, puntos y auditoría antes de decidir"),
        ("Actuar", "Usar sólo la acción disponible para el rol y registrar una razón clara"),
    ])

    add_section_page(doc, "Supervisor", "8. Revisión inicial y validaciones", "La primera decisión confirma si la idea es pertinente, suficientemente clara y qué áreas de soporte deben intervenir.")
    doc.add_heading("Revisión del supervisor", level=2)
    add_steps(doc, [
        ("Abrir pendientes", "Ingresar a Mi trabajo o a la bandeja asignada al área"),
        ("Verificar contexto", "Confirmar área, problema, propuesta, beneficio e impacto"),
        ("Elegir decisión", "Aprobar, solicitar información o rechazar"),
        ("Justificar", "Escribir una explicación concreta; evitar respuestas de una sola palabra"),
        ("Confirmar impactos", "Indicar si requiere Calidad, Seguridad y/o Mantenimiento"),
        ("Guardar", "Verificar que el folio cambie de etapa y que la siguiente bandeja reciba la tarea"),
    ])
    doc.add_heading("Criterio para áreas de soporte", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2100, 3630, 3630])
    for i, text in enumerate(("Área", "Revisar", "Registrar")):
        set_cell_shading(table.rows[0].cells[i], RED)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), 9.5, True, WHITE)
    set_repeat_table_header(table.rows[0])
    support_rows = [
        ("Calidad / Inocuidad", "Producto, contaminación, higiene, especificación, trazabilidad", "Riesgo, controles requeridos y decisión"),
        ("Seguridad", "Peligro, bloqueo, EPP, ergonomía, tránsito y cumplimiento", "Condiciones, controles y decisión"),
        ("Mantenimiento", "Diseño, instalación, disponibilidad, costo, paro y refacciones", "Factibilidad, restricciones y decisión"),
    ]
    for idx, row in enumerate(support_rows):
        cells = table.add_row().cells
        if idx % 2:
            for c in cells:
                set_cell_shading(c, LIGHT)
        for j, text in enumerate(row):
            set_run(cells[j].paragraphs[0].add_run(text), 9.5, j == 0, INK)
    set_table_geometry(table, [2100, 3630, 3630])
    add_callout(doc, "Validaciones paralelas", "Cada área decide sobre su especialidad. PROpEx no debe avanzar hasta que todas las validaciones obligatorias estén resueltas.", fill="EDF5FC", accent=BLUE)

    add_section_page(doc, "Mejora Continua", "9. Clasificar, priorizar y asignar", "Después de aprobar las validaciones, Mejora Continua transforma la idea en un trabajo ejecutable y medible.")
    add_figure(doc, "03-panel-principal.png", "Figura 7. Mi trabajo muestra pendientes que requieren clasificación o asignación.")
    doc.add_heading("Procedimiento", level=2)
    add_steps(doc, [
        ("Revisar expediente", "Confirmar que las decisiones y controles estén completos"),
        ("Clasificar", "Definir si corresponde a acción rápida, idea de mejora, proyecto Kaizen u otro tipo vigente"),
        ("Priorizar", "Evaluar riesgo, urgencia, valor, esfuerzo y dependencias"),
        ("Asignar", "Elegir un responsable con capacidad real de ejecución"),
        ("Comprometer fecha", "Definir una fecha acordada, no una fecha unilateral"),
        ("Definir resultado", "Establecer qué evidencia demostrará la implementación y el beneficio"),
        ("Escalar si aplica", "Crear un desborde cuando la decisión exceda autoridad, presupuesto o capacidad"),
    ])
    add_callout(doc, "Control", "Una idea no debe permanecer sin responsable ni fecha una vez autorizada para implementar.", fill="FFF7E6", accent=AMBER)

    add_section_page(doc, "Responsable de implementación", "10. Ejecutar, evidenciar y cerrar", "El responsable mantiene actualizado el avance. Mejora Continua valida el cierre y el sistema conserva la evidencia.")
    doc.add_heading("Registro de avances", level=2)
    add_steps(doc, [
        ("Abrir el folio", "Confirmar alcance, fecha, controles y responsable"),
        ("Registrar avance", "Describir actividad realizada, fecha, resultado parcial y siguiente paso"),
        ("Reportar bloqueo", "Indicar causa, ayuda requerida e impacto en la fecha"),
        ("Actualizar compromiso", "Solicitar ajuste cuando exista una causa válida y documentarla"),
        ("Cargar evidencia", "Adjuntar antes/después, documento o dato que demuestre el resultado"),
        ("Solicitar validación", "Marcar la implementación como lista para revisión final"),
    ])
    doc.add_heading("Cierre por Mejora Continua", level=2)
    add_bullets(doc, [
        "Comprobar que la acción implementada corresponde al alcance aprobado.",
        "Validar evidencia y resultado contra el beneficio esperado.",
        "Confirmar que no queden riesgos o acciones abiertas sin responsable.",
        "Aplicar las reglas de puntos vigentes y registrar la justificación.",
        "Cerrar, notificar y verificar actualización de tablero y reportes.",
    ])
    add_callout(doc, "No cerrar cuando", "Falta evidencia, el resultado no fue validado, existen acciones críticas abiertas o la solución introduce un riesgo no controlado.", fill="FDECEC", accent=RED_DARK)

    add_section_page(doc, "Módulos complementarios", "11. Proyectos Kaizen y recorridos Genba", "Las ideas de mayor alcance pueden convertirse en proyectos; los recorridos Genba generan hallazgos y actividades con seguimiento propio.")
    add_figure(doc, "05-proyectos-kaizen.png", "Figura 8. Portafolio de proyectos Kaizen con avance y ahorro estimado.")
    doc.add_heading("Proyecto Kaizen", level=2)
    add_bullets(doc, [
        "Definir charter, problema, alcance, objetivo, líder, equipo e indicador base.",
        "Desglosar actividades con responsable, fecha y estatus.",
        "Actualizar el porcentaje con base en actividades completadas, no por percepción.",
        "Documentar beneficios con método, periodo y fuente del cálculo.",
    ])
    doc.add_heading("Recorrido Genba", level=2)
    add_bullets(doc, [
        "Registrar área, fecha, coordinador, asistentes y propósito del recorrido.",
        "Convertir cada hallazgo en una acción clara con responsable y fecha.",
        "Cerrar hallazgos sólo cuando exista verificación en piso.",
        "Escalar restricciones estructurales mediante el mecanismo de desborde.",
    ])

    add_section_page(doc, "Gestión visual", "12. Tablero, reportes y notificaciones", "Estas vistas resumen la misma información del folio y apoyan la gestión diaria y la toma de decisiones.")
    add_figure(doc, "07-tablero-directivo.png", "Figura 9. Tablero directivo con volumen, avance, puntos, ahorro y distribución.")
    add_figure(doc, "09-reportes.png", "Figura 10. Exportación consolidada a Excel.")
    doc.add_heading("Uso recomendado", level=2)
    add_bullets(doc, [
        "Tablero: revisar tendencias, acumulación por estatus, cierre y valor; investigar el folio antes de concluir.",
        "Reportes: descargar el Excel para análisis autorizado; no usarlo como base paralela de captura.",
        "Notificaciones: atender el aviso desde la bandeja y confirmar la acción dentro del folio.",
        "Auditoría: consultar quién cambió qué y cuándo; no alterar la historia para corregir un error.",
    ])
    add_callout(doc, "Canal oficial", "Si el correo o Teams no están configurados, la notificación permanece en la bandeja local. Esto no elimina la obligación de revisar Mi trabajo.", fill="EDF5FC", accent=BLUE)

    add_section_page(doc, "Administración y control", "13. Reglas operativas y lista de verificación", "La confiabilidad de PROpEx depende de accesos correctos, datos completos y una rutina de revisión.")
    doc.add_heading("Administrador", level=2)
    add_bullets(doc, [
        "Mantener usuarios activos, roles, accesos, áreas y supervisores actualizados.",
        "Revisar reglas de puntos antes de modificarlas y comunicar la vigencia.",
        "Probar QR, exportación, almacenamiento de evidencias y notificaciones después de cambios.",
        "Restringir privilegios al mínimo necesario y retirar accesos al cambiar de función.",
        "Conservar respaldos y revisar periódicamente la bitácora de auditoría.",
    ])
    doc.add_heading("Rutina diaria", level=2)
    checklist = [
        "Cada rol revisó Mi trabajo y sus notificaciones.",
        "No existen ideas aprobadas sin responsable o fecha.",
        "Las solicitudes de información tienen respuesta o seguimiento.",
        "Los bloqueos críticos están escalados.",
        "Los avances y evidencias recientes están registrados en el folio.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        set_run(p.add_run("☐  "), 12, False, RED)
        set_run(p.add_run(item), 10.5, False, INK)
    doc.add_heading("Criterios mínimos de aceptación", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2650, 4610, 2100])
    for i, text in enumerate(("Etapa", "Debe existir", "Resultado")):
        set_cell_shading(table.rows[0].cells[i], RED)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), 9.5, True, WHITE)
    set_repeat_table_header(table.rows[0])
    acceptance = [
        ("Captura", "Problema, propuesta, beneficio y área", "Folio creado"),
        ("Supervisor", "Decisión justificada e impactos", "Ruta definida"),
        ("Validaciones", "Todas las obligatorias resueltas", "Autorización o rechazo"),
        ("Plan", "Clasificación, prioridad, responsable y fecha", "Trabajo ejecutable"),
        ("Implementación", "Avances, bloqueos y evidencia", "Listo para validar"),
        ("Cierre", "Resultado confirmado, puntos y auditoría", "Folio cerrado"),
    ]
    for idx, row in enumerate(acceptance):
        cells = table.add_row().cells
        if idx % 2:
            for c in cells:
                set_cell_shading(c, LIGHT)
        for j, text in enumerate(row):
            set_run(cells[j].paragraphs[0].add_run(text), 9.5, j == 0, INK)
    set_table_geometry(table, [2650, 4610, 2100])

    doc.add_page_break()
    add_kicker(doc, "Soporte y actualización")
    doc.add_heading("14. Control del documento", level=1)
    paragraph(doc, "Este manual debe revisarse cuando cambien pantallas, roles, reglas de negocio o responsabilidades. La versión liberada debe estar disponible para todos los usuarios.", 11.5, color=MUTED)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2600, 6760])
    control = [
        ("Propietario del proceso", "Mejora Continua"),
        ("Aprobación operativa", "Por definir"),
        ("Próxima revisión", "Al liberar cambios relevantes o en un máximo de 12 meses"),
        ("Canal de soporte", "Definir correo, extensión o responsable interno"),
    ]
    for idx, (label, value) in enumerate(control):
        set_cell_shading(table.rows[idx].cells[0], PALE_RED)
        set_run(table.rows[idx].cells[0].paragraphs[0].add_run(label), 10, True, RED_DARK)
        set_run(table.rows[idx].cells[1].paragraphs[0].add_run(value), 10, False, INK)
    set_table_geometry(table, [2600, 6760])
    doc.add_heading("Historial", level=2)
    hist = doc.add_table(rows=2, cols=4)
    set_table_geometry(hist, [1500, 1700, 4300, 1860])
    for i, h in enumerate(("Versión", "Fecha", "Cambio", "Responsable")):
        set_cell_shading(hist.rows[0].cells[i], RED)
        set_run(hist.rows[0].cells[i].paragraphs[0].add_run(h), 9.5, True, WHITE)
    values = ("1.0", "13 jul 2026", "Creación del manual integral y flujos de información", "Mejora Continua")
    for i, v in enumerate(values):
        set_run(hist.rows[1].cells[i].paragraphs[0].add_run(v), 9.5, False, INK)
    set_table_geometry(hist, [1500, 1700, 4300, 1860])
    add_callout(doc, "Validación recomendada", "Realizar una sesión breve con un operador, un supervisor, cada área validadora, un responsable y Mejora Continua. Ajustar ejemplos y nombres de campos antes de declarar esta versión como controlada.")

    # Apply header/footer to all sections and save.
    for section in doc.sections:
        configure_header_footer(section)
    doc.core_properties.title = "Manual de usuario e instructivo de trabajo PROpEx"
    doc.core_properties.subject = "Uso del sistema PROpEx por perfil y flujo de información"
    doc.core_properties.author = "Proboca - Mejora Continua"
    doc.core_properties.keywords = "PROpEx, Proboca, manual, instructivo, ideas de mejora, Kaizen, Genba"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
